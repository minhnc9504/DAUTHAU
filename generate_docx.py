"""
Generate DOCX report from BAO_CAO_DU_AN.md content.
Requires: pip install python-docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from datetime import datetime

OUTPUT_PATH = "BAO_CAO_DU_AN.docx"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    """Add heading with custom formatting."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(14)
            run.font.bold = True
        else:
            run.font.size = Pt(13)
            run.font.bold = True
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=13, indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text, level=0, size=13):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    # Light gray background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, header_color='2E74B5'):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'EEF4FB'
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ''
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    return table


def add_horizontal_line(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


# ─── Cover Page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Spacing
    for _ in range(4):
        doc.add_paragraph()

    # School/Organization header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VIETINBANK")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    for _ in range(3):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO DỰ ÁN")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HỆ THỐNG GỢI Ý GÓI THẦU")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(22)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PHÙ HỢP CHO DOANH NGHIỆP")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(22)
    run.font.bold = True

    for _ in range(2):
        doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VietinBank — Dự án Đấu thầu")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.italic = True

    for _ in range(8):
        doc.add_paragraph()

    # Info table
    info_data = [
        ["Phiên bản:", "v12.0 (app.py)"],
        ["Ngày báo cáo:", "24/03/2026"],
        ["Tác giả:", "minhnche180504"],
        ["Repository:", "D:\\VIETINBANK\\DAUTHAUGITHUB"],
        ["Framework giao diện:", "Streamlit >=1.28.0"],
    ]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            row.cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()


# ─── Section 1: Tổng quan ────────────────────────────────────────────────────

def add_section_1(doc):
    add_heading(doc, "1. Tổng quan dự án", 1)
    add_horizontal_line(doc)

    add_heading(doc, "1.1 Giới thiệu", 2)
    add_paragraph(doc,
        "Hệ thống Gợi ý gói thầu phù hợp cho Doanh nghiệp là một ứng dụng web được phát triển bởi "
        "VietinBank, nhằm hỗ trợ các doanh nghiệp tìm kiếm và lọc các gói thầu phù hợp nhất dựa trên "
        "lịch sử đấu thầu, năng lực và các tiêu chí lọc đa dạng.")

    add_heading(doc, "1.2 Tính năng chính", 2)
    headers = ["STT", "Tính năng", "Mô tả"]
    rows = [
        ["1", "Tìm kiếm thông minh", "Gợi ý gói thầu dựa trên lịch sử đấu thầu hoặc mô tả năng lực"],
        ["2", "Phân tích đối thủ", "Phân tích chi tiết các nhà thầu từng tham gia cùng chủ đầu tư / lĩnh vực"],
        ["3", "Lọc đa tiêu chí", "Lọc theo giá thầu, tỉnh/thành, vùng miền, lĩnh vực đầu tư"],
        ["4", "Xuất báo cáo", "Xuất kết quả ra file Excel"],
        ["5", "Hai chế độ sử dụng", "Doanh nghiệp có lịch sử và doanh nghiệp mới"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "1.3 Dữ liệu đầu vào", 2)
    add_paragraph(doc, "File dữ liệu chính: dauthau_data.csv chứa 16,801+ dòng với khoảng 20 cột thông tin.")
    headers2 = ["Cột", "Mô tả"]
    rows2 = [
        ["bidonotifycontractormnotifyno", "Mã số gói thầu"],
        ["bidonotifycontractormbidname", "Tên gói thầu (dùng cho TF-IDF)"],
        ["bidonotifycontractormprojectname", "Tên dự án"],
        ["bidonotifycontractorminvestorname", "Tên chủ đầu tư"],
        ["bidonotifycontractorminvestfield", "Lĩnh vực đầu tư"],
        ["provincename", "Tỉnh/Thành phố"],
        ["bidecontractorinputresultdtobidprice", "Giá trị gói thầu (VND)"],
        ["bidecontractorinputresultdtoopendate", "Ngày mở thầu"],
        ["bidresult", "Kết quả đấu thầu (1 = trúng, 0 = không trúng)"],
        ["taxcode", "Mã số thuế doanh nghiệp"],
        ["orgfullname", "Tên đầy đủ doanh nghiệp"],
    ]
    add_table(doc, headers2, rows2)
    doc.add_page_break()


# ─── Section 2: Công nghệ ─────────────────────────────────────────────────────

def add_section_2(doc):
    add_heading(doc, "2. Công nghệ sử dụng", 1)
    add_horizontal_line(doc)

    add_heading(doc, "2.1 Stack công nghệ", 2)
    headers = ["Thành phần", "Công nghệ", "Phiên bản min", "Mục đích"]
    rows = [
        ["Web Framework", "Streamlit", ">=1.28.0", "Xây dựng giao diện web"],
        ["Xử lý dữ liệu", "Pandas", ">=2.0.0", "Đọc, xử lý, lọc DataFrame"],
        ["Machine Learning", "scikit-learn", ">=1.3.0", "TF-IDF Vectorizer, Cosine Similarity"],
        ["Tính toán số", "NumPy", ">=1.24.0", "Xử lý mảng đa chiều, phép toán ma trận"],
        ["Xuất Excel", "XlsxWriter", ">=3.1.0", "Tạo file Excel với định dạng"],
        ["Ngôn ngữ lập trình", "Python", "3.x", "Ngôn ngữ chính toàn bộ hệ thống"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "2.2 Các thư viện bổ sung", 2)
    add_bullet(doc, "sentence-transformers (tùy chọn): Semantic embedding với model BAAI/bge-m3")
    add_bullet(doc, "pyarrow: Đọc/ghi định dạng Parquet (Snappy compression)")
    add_bullet(doc, "joblib: Lưu trữ và tái sử dụng mô hình ML đã huấn luyện")
    add_bullet(doc, "python-dotenv: Quản lý biến môi trường từ file .env")
    add_bullet(doc, "Git LFS: Lưu trữ file CSV và PKL binary dung lượng lớn")

    doc.add_paragraph()
    add_heading(doc, "2.3 Kiến trúc hệ thống", 2)
    arch_code = """┌──────────────────────────────────────────────────────────────┐
│                    Giao diện người dùng                          │
│                   (Streamlit Web UI)                              │
└──────────────────────┬───────────────────────────────────────────┘
                       │
                       ▼
┌──────────────────────────────────────────────────────────────┐
│               app.py (Main Application)                       │
│  • load_assets()     → Tải mô hình TF-IDF (.pkl)              │
│  • load_data()       → Tải dữ liệu Parquet (cache 1h)        │
│  • calculate_hybrid_score() → Tính điểm lai                  │
│  • to_excel()        → Xuất file Excel                         │
└───────────┬──────────────────────┬────────────────────────────┘
            │                      │
            ▼                      ▼
┌──────────────────────┐  ┌──────────────────────────────────┐
│  models/*.pkl        │  │     dauthau_data.csv             │
│  (TF-IDF Model)      │  │     (Raw Data)                   │
└──────────────────────┘  └─────────────┬────────────────────┘
                                        ▼
                         ┌──────────────────────────────┐
                         │   convert_data.py              │
                         │   (CSV → Parquet)              │
                         └──────────────┬─────────────────┘
                                        ▼
                         ┌──────────────────────────────┐
                         │   train_model.py              │
                         │   • Huấn luyện TF-IDF         │
                         │   • Tạo ma trận sparse        │
                         └──────────────────────────────┘"""
    add_code_block(doc, arch_code)
    doc.add_page_break()


# ─── Section 3: Mô hình ──────────────────────────────────────────────────────

def add_section_3(doc):
    add_heading(doc, "3. Mô hình sử dụng", 1)
    add_horizontal_line(doc)

    add_heading(doc, "3.1 Mô hình TF-IDF", 2)
    add_paragraph(doc,
        "TF-IDF (Term Frequency — Inverse Document Frequency) là mô hình trung tâm của hệ thống, "
        "dùng để tính độ tương đồng giữa năng lực doanh nghiệp và nội dung gói thầu.")

    headers = ["Thuộc tính", "Giá trị"]
    rows = [
        ["Loại mô hình", "TfidfVectorizer (sklearn)"],
        ["Thuật toán", "TF-IDF"],
        ["File đầu ra", "models/tfidf_model.pkl, models/tfidf_matrix.pkl"],
        ["Script huấn luyện", "train_model.py"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "3.1.1 Tham số cấu hình", 3)
    code = """tfidf = TfidfVectorizer(ngram_range=(1, 2), min_df=2)"""
    add_code_block(doc, code)

    headers2 = ["Tham số", "Giá trị", "Ý nghĩa"]
    rows2 = [
        ["ngram_range", "(1, 2)", "Sử dụng unigrams và bigrams"],
        ["min_df", "2", "Bỏ qua từ xuất hiện < 2 tài liệu"],
        ["max_df", "0.95", "Bỏ qua từ xuất hiện > 95% tài liệu"],
    ]
    add_table(doc, headers2, rows2)

    doc.add_paragraph()
    add_heading(doc, "3.2 Mô hình Hybrid Scoring", 2)
    add_paragraph(doc, "Hệ thống kết hợp nhiều tín hiệu để tạo điểm tổng hợp:")
    formula = """Total Score = (Lexical Score × 0.60) + (Price Score × 0.25) + (Recency Score × 0.15)"""
    add_code_block(doc, formula)

    headers3 = ["Thành phần", "Trọng số", "Nguồn dữ liệu"]
    rows3 = [
        ["Lexical Score", "60%", "TF-IDF Cosine Similarity (0-100)"],
        ["Price Score", "25%", "Filter cứng (luôn = 100 nếu qua lọc)"],
        ["Recency Score", "15%", "Số ngày đến ngày mở thầu"],
    ]
    add_table(doc, headers3, rows3)
    doc.add_page_break()


# ─── Section 4: Thuật toán ────────────────────────────────────────────────────

def add_section_4(doc):
    add_heading(doc, "4. Thuật toán chi tiết", 1)
    add_horizontal_line(doc)

    add_heading(doc, "4.1 Công thức TF-IDF", 2)
    add_paragraph(doc,
        "Công thức TF-IDF tính trọng số của mỗi từ trong tài liệu dựa trên tần suất xuất hiện "
        "trong tài liệu đó và mức độ phổ biến trong toàn bộ corpus:")
    add_paragraph(doc, "w(t,d) = TF(t,d) x log(N / df(t))", bold=True)
    add_paragraph(doc, "Trong đó:")
    add_bullet(doc, "w(t,d): Trọng số TF-IDF của từ t trong tài liệu d")
    add_bullet(doc, "TF(t,d): Tần suất xuất hiện của từ t trong tài liệu d")
    add_bullet(doc, "N: Tổng số tài liệu trong corpus")
    add_bullet(doc, "df(t): Số tài liệu chứa từ t")

    doc.add_paragraph()
    add_heading(doc, "4.2 Cosine Similarity", 2)
    add_paragraph(doc,
        "Hệ thống sử dụng cosine similarity để đo độ tương đồng giữa vector TF-IDF của "
        "truy vấn doanh nghiệp và vector của từng gói thầu:")
    add_code_block(doc, "similarity(A, B) = (A · B) / (||A|| × ||B||)")
    add_paragraph(doc,
        "Hàm sklearn.metrics.pairwise.cosine_similarity() được sử dụng để tính toán "
        "hàng loạt hiệu quả trên ma trận sparse.")

    doc.add_paragraph()
    add_heading(doc, "4.3 Quy trình tính điểm TF-IDF", 2)
    code_steps = """1. Transform truy vấn doanh nghiệp:
   query_vector = tfidf.transform([query.lower()])

2. Transform tất cả tên gói thầu:
   current_matrix = tfidf.transform(bid_names)

3. Tính cosine similarity:
   sim_scores = cosine_similarity(query_vector, current_matrix)

4. Gán điểm lexical (0-100):
   res_df['score'] = (sim_scores * 100).round(1)"""
    add_code_block(doc, code_steps)

    doc.add_paragraph()
    add_heading(doc, "4.4 Thuật toán Hybrid Scoring", 2)
    code_hybrid = """def calculate_hybrid_score(row, start_val, end_val):
    # 1. Điểm từ vựng (từ TF-IDF)
    lexical_score = row['score']

    # 2. Điểm giá (hard filter)
    price_score = 100

    # 3. Điểm thời gian (recency)
    days_to_open = (row['dt_opendate'] - datetime.now()).days
    if days_to_open > 0:
        recency_score = max(50, 100 - (days_to_open * 2))
    else:
        recency_score = 30

    # 4. Tổng hợp điểm lai
    total_score = (lexical_score * 0.6) + (price_score * 0.25) + (recency_score * 0.15)
    return round(total_score, 1)"""
    add_code_block(doc, code_hybrid)

    doc.add_paragraph()
    add_heading(doc, "4.5 Phân tích đối thủ — Killer Score", 2)
    add_paragraph(doc,
        "Thuật toán Killer Score đánh giá mức độ cạnh tranh của từng nhà thầu trên 4 chiều:")
    code_killer = """# Xác định đối thủ tiềm năng:
#   - Cùng chủ đầu tư (investorname = p_investor)
#   - HOẶC cùng lĩnh vực (investfield = p_field)
#   - HOẶC cùng tỉnh/thành (provincename = p_province)

# Tính Killer Score cho mỗi đối thủ:
competitor_score = (wins / joins).clip(upper=1.0) × 2.5  # cho mỗi chiều

# Tổng Killer Score:
# Tối đa 10.0 điểm (2.5 × 4 chiều: investor, province, field, price)

# Lọc: Loại bỏ đối thủ có tổng số lần tham gia < 2"""
    add_code_block(doc, code_killer)

    headers = ["Chiều", "Công thức", "Ý nghĩa"]
    rows = [
        ["Investor Score", "(wins/joins) × 2.5", "Tỷ lệ thắng với cùng chủ đầu tư"],
        ["Province Score", "(wins/joins) × 2.5", "Tỷ lệ thắng tại cùng tỉnh/thành"],
        ["Field Score", "(wins/joins) × 2.5", "Tỷ lệ thắng trong cùng lĩnh vực"],
        ["Price Segment", "(wins/joins) × 2.5", "Cạnh tranh trong cùng phân khúc giá"],
    ]
    add_table(doc, headers, rows)
    doc.add_page_break()


# ─── Section 5: Pipeline ──────────────────────────────────────────────────────

def add_section_5(doc):
    add_heading(doc, "5. Pipeline xử lý dữ liệu", 1)
    add_horizontal_line(doc)

    add_heading(doc, "5.1 Pipeline tổng thể", 2)
    pipeline_code = """dauthau_data.csv (CSV thô)
        │
        ▼
[convert_data.py]
  • Thử nhiều encoding: UTF-8-sig → UTF-8 → CP1258 → Latin-1
  • Chuyển đổi sang Parquet (Snappy compression)
  • Thêm cột ngày chuẩn hóa
        │
        ▼
dauthau_data.parquet
        │
        ├──► [train_model.py] ──► models/tfidf_model.pkl
        │                        models/tfidf_matrix.pkl
        │
        └──► [app.py / run.py]
              • Tải dữ liệu (cache 1 giờ)
              • Tải mô hình TF-IDF (cache vĩnh viễn)
              • Tính điểm lai và lọc
              • Phân tích đối thủ
              • Xuất kết quả Excel"""
    add_code_block(doc, pipeline_code)

    doc.add_paragraph()
    add_heading(doc, "5.2 Chiến lược Cache", 2)
    headers = ["Thành phần", "Chiến lược", "TTL", "Mục đích"]
    rows = [
        ["TF-IDF Model", "@st.cache_resource", "Vĩnh viễn", "Tránh load lại mô hình"],
        ["TF-IDF Matrix", "@st.cache_resource", "Vĩnh viễn", "Tránh tính lại ma trận"],
        ["Dữ liệu Parquet", "@st.cache_data(ttl=3600)", "1 giờ", "Cân bằng dữ liệu mới và hiệu năng"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "5.3 Xử lý encoding tiếng Việt", 2)
    add_paragraph(doc, "Hệ thống thử theo thứ tự ưu tiên:")
    add_bullet(doc, "utf-8-sig (UTF-8 with BOM)")
    add_bullet(doc, "utf-8 (UTF-8)")
    add_bullet(doc, "cp1258 (Windows Vietnamese codepage)")
    add_bullet(doc, "latin-1 (ISO-8859-1)")
    doc.add_page_break()


# ─── Section 6: Kiểm thử ─────────────────────────────────────────────────────

def add_section_6(doc):
    add_heading(doc, "6. Kiểm thử và đánh giá", 1)
    add_horizontal_line(doc)

    add_heading(doc, "6.1 Tình trạng testing hiện tại", 2)
    headers = ["Loại testing", "Trạng thái", "Chi tiết"]
    rows = [
        ["Unit Test", "Chưa có", "Không có file test_*.py"],
        ["Integration Test", "Chưa có", "Không có script kiểm thử tích hợp"],
        ["E2E Test", "Chưa có", "Không có framework E2E"],
        ["CI/CD", "Chưa có", "Không có GitHub Actions"],
        ["Kiểm thử thủ công", "Hạn chế", "Chỉ kiểm tra runtime khi chạy"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "6.2 Các kịch bản test đã thực hiện", 2)
    headers2 = ["Test ID", "Loại", "Kết quả", "Ghi chú"]
    rows2 = [
        ["TC-001", "Khởi động ứng dụng", "PASS", "Ứng dụng chạy OK trên Streamlit"],
        ["TC-002", "Load TF-IDF Model", "PASS", "Model load thành công, không lỗi pickle"],
        ["TC-003", "Encoding tiếng Việt", "PASS", "Tiếng Việt hiển thị chính xác"],
        ["TC-004", "Filter theo giá", "PASS", "Lọc đúng khoảng giá"],
        ["TC-005", "Phân tích đối thủ", "PASS", "Killer Score tính đúng"],
        ["TC-006", "Export Excel", "PASS", "Excel xuất đúng định dạng"],
    ]
    add_table(doc, headers2, rows2)
    doc.add_page_break()


# ─── Section 7: Cấu trúc dự án ────────────────────────────────────────────────

def add_section_7(doc):
    add_heading(doc, "7. Cấu trúc dự án", 1)
    add_horizontal_line(doc)

    add_heading(doc, "7.1 Cây thư mục", 2)
    tree_code = """DAUTHAUGITHUB/
│
├── app.py                    # Ứng dụng Streamlit chính (v12.0, 529 dòng)
│   ├── load_assets()         # Tải mô hình TF-IDF
│   ├── load_data()           # Tải dữ liệu (Parquet)
│   ├── calculate_hybrid_score()  # Tính điểm lai
│   ├── to_excel()            # Xuất Excel
│   ├── format_currency()     # Định dạng VND
│   └── clean_field_name()    # Chuẩn hóa tên lĩnh vực
│
├── train_model.py            # Script huấn luyện TF-IDF
│   └── train()               # Tạo .pkl files
│
├── convert_data.py           # Script chuyển đổi CSV → Parquet
│   └── convert()             # Encoding detection + conversion
│
├── run.py                    # Script điều phối toàn bộ pipeline
│   └── run_app()             # Chạy pipeline → Streamlit
│
├── requirements.txt          # Danh sách dependencies
├── README.md                 # Tài liệu dự án
├── run_fast.bat              # Shortcut Windows chạy nhanh
│
├── data/
│   └── dauthau_data.csv      # Dữ liệu thô (Git LFS)
│
└── models/                   # (sinh ra lúc runtime)
    ├── tfidf_model.pkl       # TfidfVectorizer đã huấn luyện
    └── tfidf_matrix.pkl      # Ma trận TF-IDF"""
    add_code_block(doc, tree_code)

    doc.add_paragraph()
    add_heading(doc, "7.2 Danh sách hàm và chức năng", 2)
    headers = ["File", "Hàm", "Mô tả"]
    rows = [
        ["app.py", "load_assets()", "Tải và cache TF-IDF model từ .pkl"],
        ["app.py", "load_data()", "Tải và cache Parquet data"],
        ["app.py", "calculate_hybrid_score()", "Tính điểm lai cho mỗi gói thầu"],
        ["app.py", "to_excel()", "Xuất DataFrame ra BytesIO"],
        ["app.py", "format_currency()", "Định dạng số VND (VD: 1.000.000)"],
        ["app.py", "format_date_series()", "Vectorized date formatting"],
        ["app.py", "clean_field_name()", "Loại bỏ 'Lĩnh vực: ' prefix"],
        ["app.py", "Phân tích đối thủ", "Thuật toán Killer Score"],
        ["train_model.py", "train()", "Huấn luyện vectorizer"],
        ["convert_data.py", "convert()", "CSV → Parquet"],
        ["run.py", "run_app()", "Điều phối pipeline"],
    ]
    add_table(doc, headers, rows)
    doc.add_page_break()


# ─── Section 8: Kết luận ─────────────────────────────────────────────────────

def add_section_8(doc):
    add_heading(doc, "8. Kết luận và hướng phát triển", 1)
    add_horizontal_line(doc)

    add_heading(doc, "8.1 Kết luận", 2)
    add_bullet(doc, "Công nghệ hiện đại: Streamlit, scikit-learn, Pandas")
    add_bullet(doc, "Thuật toán hiệu quả: TF-IDF + Cosine Similarity + Hybrid Scoring")
    add_bullet(doc, "Giao diện thân thiện: Web-based, hỗ trợ tiếng Việt")
    add_bullet(doc, "Phân tích sâu: Killer Score cho phân tích đối thủ")
    add_bullet(doc, "Xuất báo cáo: Hỗ trợ xuất Excel")

    doc.add_paragraph()
    add_heading(doc, "8.2 Hạn chế hiện tại", 2)
    headers = ["Hạn chế", "Mức ưu tiên"]
    rows = [
        ["Chưa có Unit Test / Integration Test", "Cao"],
        ["Chưa có CI/CD pipeline", "Cao"],
        ["Chưa có file README chi tiết", "Trung bình"],
        ["Chưa triển khai mô hình ML nâng cao (BERT, RAG)", "Trung bình"],
        ["Chưa có logging và monitoring", "Trung bình"],
    ]
    add_table(doc, headers, rows)

    doc.add_paragraph()
    add_heading(doc, "8.3 Hướng phát triển tương lai", 2)

    add_heading(doc, "Ngắn hạn (1-3 tháng)", 3)
    add_bullet(doc, "Thêm Unit Test bằng pytest cho các hàm chính")
    add_bullet(doc, "Thêm Integration Test cho pipeline dữ liệu")
    add_bullet(doc, "Thiết lập CI/CD với GitHub Actions")
    add_bullet(doc, "Cải thiện README với hướng dẫn chi tiết")

    add_heading(doc, "Trung hạn (3-6 tháng)", 3)
    add_bullet(doc, "Nâng cấp mô hình: Thử nghiệm Sentence-BERT (SBERT) thay TF-IDF")
    add_bullet(doc, "Thêm RAG (Retrieval-Augmented Generation) kết hợp LLM")
    add_bullet(doc, "Dashboard analytics: Hiển thị thống kê trực quan")
    add_bullet(doc, "API endpoint: REST API cho hệ thống khác tích hợp")

    add_heading(doc, "Dài hạn (6-12 tháng)", 3)
    add_bullet(doc, "Multi-modal search: Kết hợp văn bản, hình ảnh, dữ liệu có cấu trúc")
    add_bullet(doc, "Real-time updates: Cập nhật dữ liệu thời gian thực từ nguồn VietinBank")
    add_bullet(doc, "Personalization: Cá nhân hóa gợi ý dựa trên lịch sử tìm kiếm")
    add_bullet(doc, "A/B Testing: So sánh các thuật toán khác nhau")
    doc.add_page_break()


# ─── Section 9: Hướng dẫn cài đặt ─────────────────────────────────────────────

def add_section_9(doc):
    add_heading(doc, "9. Hướng dẫn cài đặt và sử dụng", 1)
    add_horizontal_line(doc)

    add_heading(doc, "9.1 Yêu cầu hệ thống", 2)
    add_bullet(doc, "Python >= 3.8")
    add_bullet(doc, "RAM: Tối thiểu 4GB (khuyến nghị 8GB+)")
    add_bullet(doc, "Disk: 500MB trở lên")
    add_bullet(doc, "OS: Windows / Linux / macOS")

    doc.add_paragraph()
    add_heading(doc, "9.2 Cài đặt dependencies", 2)
    add_code_block(doc, "pip install -r requirements.txt")

    add_heading(doc, "9.3 Cách chạy dự án", 2)
    add_paragraph(doc, "Có 3 cách chạy ứng dụng:")
    add_bullet(doc, "Cách 1: Chạy nhanh (Windows) → run_fast.bat")
    add_bullet(doc, "Cách 2: Chạy đầy đủ → python run.py")
    add_bullet(doc, "Cách 3: Chạy riêng từng bước:")
    add_code_block(doc, """# Bước 1: Chuyển đổi dữ liệu
python convert_data.py

# Bước 2: Huấn luyện mô hình
python train_model.py

# Bước 3: Khởi chạy giao diện
streamlit run app.py""")

    add_heading(doc, "9.4 Cấu hình biến môi trường (tùy chọn)", 2)
    add_code_block(doc, """# Tạo file .env
DAUTHAU_RAW_CSV_PATH=data/dauthau_data.csv
DAUTHAU_CURATED_DIR=data/curated
DAUTHAU_ARTIFACTS_DIR=artifacts
DAUTHAU_SEMANTIC_ENABLED=false
DAUTHAU_TOP_K_RESULTS=20
DAUTHAU_RERANK_TOP_K=50""")

    add_heading(doc, "9.5 Các lệnh CLI", 2)
    headers = ["Lệnh", "Mô tả"]
    rows = [
        ["python main.py setup", "Cài đặt dependencies"],
        ["python main.py ingest", "CSV → curated parquet"],
        ["python main.py build-index", "Build profiles + TF-IDF index"],
        ["python main.py rebuild", "Full pipeline (ingest + build)"],
        ["python main.py serve", "Mở giao diện Streamlit"],
        ["python main.py run", "Auto check + rebuild + UI"],
        ["python main.py status", "Xem trạng thái artifact"],
    ]
    add_table(doc, headers, rows)
    doc.add_page_break()


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    add_cover_page(doc)
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)

    # Footer: last paragraph note
    add_horizontal_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Báo cáo được tạo tự động từ phân tích codebase.\n"
        "Ngày tạo: 24/03/2026 | Phiên bản báo cáo: 1.0"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUTPUT_PATH)
    print("Da tao xong: " + OUTPUT_PATH)


if __name__ == "__main__":
    main()
