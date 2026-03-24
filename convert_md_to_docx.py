"""
Chuyển đổi GIAI_THICH_CODE.md sang GIAI_THICH_CODE.docx
Sử dụng: python convert_md_to_docx.py
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

def hex_to_rgb(hex_str: str):
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    """Đặt màu nền cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Đặt border cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        if edge in kwargs:
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
        else:
            tag.set(qn('w:val'), 'single')
            tag.set(qn('w:sz'), '4')
            tag.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


# ──────────────────────────────────────────────────────────────────────────────
# Parser
# ──────────────────────────────────────────────────────────────────────────────

def parse_md(content: str):
    """
    Parse markdown text thành danh sách block.
    Block types: heading, table, code_block, hr, list, paragraph
    """
    lines = content.split('\n')
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            blocks.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # Table header (| ... |)
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse table
            rows = []
            for tline in table_lines:
                # Skip separator line (|---|---|)
                if re.match(r'^\|[\s\-:|]+\|$', tline):
                    continue
                cols = [c.strip() for c in tline.strip('|').split('|')]
                rows.append(cols)
            if rows:
                blocks.append({'type': 'table', 'rows': rows})
            continue

        # Code block
        if line.startswith('```'):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({'type': 'code_block', 'lang': lang, 'lines': code_lines})
            continue

        # HR
        if line.strip() in ('---', '***', '___'):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # List
        if re.match(r'^\s*[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                items.append(re.sub(r'^\s*[-*+]\s+', '', lines[i]).strip())
                i += 1
            blocks.append({'type': 'list', 'items': items})
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\s*\d+\.\s+', '', lines[i]).strip())
                i += 1
            blocks.append({'type': 'ordered_list', 'items': items})
            continue

        # Paragraph
        if line.strip():
            para_lines = []
            while i < len(lines) and lines[i].strip() and \
                    not lines[i].startswith('#') and \
                    not lines[i].startswith('```') and \
                    not re.match(r'^\s*[-*+]\s+', lines[i]) and \
                    not re.match(r'^\s*\d+\.\s+', lines[i]) and \
                    not lines[i].strip().startswith('|'):
                para_lines.append(lines[i].rstrip())
                i += 1
            text = ' '.join(para_lines).strip()
            if text:
                blocks.append({'type': 'paragraph', 'text': text})
            continue

        i += 1

    return blocks


def inline_format(para, text, styles):
    """Parse inline **bold**, `code`, và links trong text."""
    # Regex: **bold**, `code`, [text](url)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            if 'color' in styles:
                run.font.color.rgb = RGBColor(*hex_to_rgb(styles['color']))
            if 'size' in styles:
                run.font.size = Pt(styles['size'])
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(styles.get('size', 11) - 1)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run = para.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                run.font.underline = True
        else:
            run = para.add_run(part)
            if 'color' in styles:
                run.font.color.rgb = RGBColor(*hex_to_rgb(styles['color']))
            if 'size' in styles:
                run.font.size = Pt(styles['size'])


# ──────────────────────────────────────────────────────────────────────────────
# Color palette
# ──────────────────────────────────────────────────────────────────────────────

COLORS = {
    'title_bg':   '1A3A5C',   # Navy cho header chính
    'h1_bg':      '1F4E79',   # Dark blue cho H1
    'h2_bg':      '2E75B6',   # Blue cho H2
    'h3_bg':      '5B9BD5',   # Light blue cho H3
    'h4_color':   '2E4057',   # Dark cho H4
    'table_hdr':  '1F4E79',   # Table header
    'table_alt':  'EBF3FB',   # Table alternating row
    'table_brd':  '2E75B6',   # Table border
    'code_bg':    'F2F2F2',   # Code block background
    'code_bdr':   'CCCCCC',   # Code block border
    'hr_color':   '2E75B6',   # HR line
    'body':       '333333',   # Body text
    'bullet':     '2E75B6',   # Bullet color
}


# ──────────────────────────────────────────────────────────────────────────────
# Builder
# ──────────────────────────────────────────────────────────────────────────────

def build_docx(blocks, output_path):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default body style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = Pt(16)

    def add_hr(doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), COLORS['hr_color'])
        pBdr.append(bottom)
        pPr.append(pBdr)

    for block in blocks:
        btype = block['type']

        # ── H1 ──────────────────────────────────────────────────────────────
        if btype == 'heading' and block['level'] == 1:
            add_hr(doc)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Background shade cho heading 1
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  COLORS['h1_bg'])
            pPr.append(shd)

        # ── H2 ──────────────────────────────────────────────────────────────
        elif btype == 'heading' and block['level'] == 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(*hex_to_rgb(COLORS['h2_bg']))

        # ── H3 ──────────────────────────────────────────────────────────────
        elif btype == 'heading' and block['level'] == 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(*hex_to_rgb(COLORS['h3_bg']))

        # ── H4 ──────────────────────────────────────────────────────────────
        elif btype == 'heading' and block['level'] == 4:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(block['text'])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(*hex_to_rgb(COLORS['h4_color']))

        # ── Paragraph ─────────────────────────────────────────────────────────
        elif btype == 'paragraph':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            inline_format(p, block['text'], {'color': COLORS['body'], 'size': 12})

        # ── Bullet list ──────────────────────────────────────────────────────
        elif btype == 'list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent  = Cm(0.5)
                p.paragraph_format.space_after  = Pt(2)
                inline_format(p, item, {'color': COLORS['body'], 'size': 12})

        # ── Numbered list ───────────────────────────────────────────────────
        elif btype == 'ordered_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                inline_format(p, item, {'color': COLORS['body'], 'size': 12})

        # ── Code block ──────────────────────────────────────────────────────
        elif btype == 'code_block':
            code_text = '\n'.join(block['lines'])
            lines_code = code_text.split('\n')
            max_len = max(len(l) for l in lines_code) if lines_code else 0

            # Determine columns based on content
            if max_len > 100:
                col_widths = (Cm(6.5), Cm(8.5))
            elif max_len > 60:
                col_widths = (Cm(4.5), Cm(10.5))
            else:
                col_widths = (Cm(3.0), Cm(12.0))

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.allow_autofit = False

            # Col 0: lang label
            cell0 = table.cell(0, 0)
            cell0.width = col_widths[0]
            p0 = cell0.paragraphs[0]
            r0 = p0.add_run(block['lang'].upper() if block['lang'] else 'CODE')
            r0.bold = True
            r0.font.size = Pt(9)
            r0.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell0, COLORS['code_bg'])

            # Col 1: code
            cell1 = table.cell(0, 1)
            cell1.width = col_widths[1]
            for idx, codeline in enumerate(lines_code):
                if idx == 0:
                    p1 = cell1.paragraphs[0]
                else:
                    p1 = cell1.add_paragraph()
                p1.paragraph_format.space_after = Pt(0)
                r = p1.add_run(codeline)
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            set_cell_bg(cell1, COLORS['code_bg'])

            # Table borders
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(cell,
                        top   = {'val': 'single', 'sz': 1, 'color': COLORS['code_bdr']},
                        bottom= {'val': 'single', 'sz': 1, 'color': COLORS['code_bdr']},
                        left  = {'val': 'single', 'sz': 1, 'color': COLORS['code_bdr']},
                        right = {'val': 'single', 'sz': 1, 'color': COLORS['code_bdr']},
                    )

            doc.add_paragraph()  # spacing after code block

        # ── Table ────────────────────────────────────────────────────────────
        elif btype == 'table':
            rows_data = block['rows']
            if not rows_data:
                continue

            num_cols = max(len(r) for r in rows_data)
            num_rows = len(rows_data)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Estimate column widths
            total_width = 16.0  # Cm available
            avg_col_w = total_width / num_cols

            for r_idx, row_data in enumerate(rows_data):
                row = table.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data):
                    cell = row.cells[c_idx]
                    cell.width = Cm(avg_col_w)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    is_header = r_idx == 0

                    if is_header:
                        run = p.add_run(cell_text)
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_bg(cell, COLORS['table_hdr'])
                    else:
                        inline_format(p, cell_text,
                                      {'color': COLORS['body'], 'size': 11})
                        # Alternating row color
                        if r_idx % 2 == 0:
                            set_cell_bg(cell, COLORS['table_alt'])

                    # Border
                    for edge in ('top', 'left', 'bottom', 'right'):
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        tag = OxmlElement(f'w:{edge}')
                        tag.set(qn('w:val'),   'single')
                        tag.set(qn('w:sz'),    '4')
                        tag.set(qn('w:color'), COLORS['table_brd'])
                        tcBorders.append(tag)
                        tcPr.append(tcBorders)

            doc.add_paragraph()

        # ── HR ───────────────────────────────────────────────────────────────
        elif btype == 'hr':
            add_hr(doc)

    doc.save(output_path)
    print(f"✓ Đã lưu: {output_path}")


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import os

    md_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'GIAI_THICH_CODE.md')
    docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              'GIAI_THICH_CODE.docx')

    print(f"Đọc: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    print("Parse markdown...")
    blocks = parse_md(content)

    print("Build DOCX...")
    build_docx(blocks, docx_path)
    print("Xong!")
